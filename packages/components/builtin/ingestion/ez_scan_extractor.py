"""LLM 驱动的文档提取组件。

从 PDF、图片、Word、Excel 或文本文件中提取文本，调用大模型转为结构化 ObservationTable。
所有可配置项通过 manifest 参数传入，代码中无硬编码。

参数（全部在 manifest 定义，网页可编辑）：
- path: 文件路径（必填），支持两种格式：
  - artifact:{artifact_id} — 从 MinIO 下载工件到临时文件再处理（推荐）
  - 文件系统路径 — 直接读取本地文件（兼容旧格式）
- prompt: LLM 提示词（包含角色设定 + 提取指令 + 输出格式要求）
- file_engine: 文件读取方式（auto / pymupdf / image / raw），默认 auto 自动检测
  - auto：根据文件后缀自动选择读取方式（推荐）
    - .pdf：先提取文字层，文字太少（平均每页 < 50 字符）自动切换到 image 模式
    - .txt/.md：直接 UTF-8 读取
    - .jpg/.jpeg/.png：返回 base64 图片列表（多模态大模型识别）
    - .doc/.docx：用 python-docx 读取文本
    - .xls/.xlsx：用 openpyxl 读取
    - 其他后缀：尝试当文本读取
  - pymupdf/image/raw：旧模式，保留向后兼容
- image_dpi: image 模式渲染分辨率
- max_content_chars: 最大内容字符数
- timeout: LLM 调用超时秒数
"""

import asyncio
import io
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Any
from uuid import UUID

import httpx

from packages.common.errors import AppError
from packages.components.builtin.types import ObservationTable
from packages.components.sdk import ComponentContext, ComponentResult


class EZScanExtractor:
    """LLM 驱动的文档提取组件。

    所有行为由参数控制，无硬编码默认值。
    """

    async def execute(
        self,
        context: ComponentContext,
        params: dict[str, Any],
    ) -> ComponentResult:
        """读取文件 → 提取文本 → 调用 LLM → 返回 ObservationTable。"""
        # 检查是否指定了专用解析工具
        tool_type: str = params.get("tool_type", "llm")
        if tool_type == "xrd_tool":
            # 委托给 XRD 确定性解析组件
            from packages.components.builtin.ingestion.xrd_tool_component import (
                XrdToolComponent,
            )

            return await XrdToolComponent().execute(context, params)

        path_str: str = params["path"]
        prompt: str = params.get("prompt", "")
        engine: str = params.get("file_engine", "auto")
        image_dpi: int = 200
        max_chars: int = 999999999
        timeout: int = 300

        if not prompt:
            raise AppError(
                code="validation_failed",
                message="缺少 prompt 参数",
                retryable=False,
            )

        # 支持 artifact:{artifact_id} 格式：从 MinIO 下载到临时文件
        # 兼容旧格式：直接使用文件系统路径
        if path_str.startswith("artifact:"):
            file_path = await self._download_artifact_to_temp(context, path_str[len("artifact:") :])
        else:
            file_path = Path(path_str)

        # 1. 提取文件文本（同步 I/O，用 to_thread 避免阻塞事件循环）
        content = await asyncio.to_thread(_extract_text, file_path, engine, image_dpi)
        if isinstance(content, str) and len(content) > max_chars:
            content = content[:max_chars]

        # 2. 判断是文本还是图片模式
        is_image_mode = isinstance(content, list)

        # 3. 空内容直接返回空表
        if (is_image_mode and len(content) == 0) or (not is_image_mode and not content.strip()):
            table = ObservationTable(
                columns=(),
                rows=(),
                source_locations=(),
            )
            return ComponentResult(
                outputs={"observations": table},
                summary="提取 0 行数据（空文件）",
                metadata={"row_count": 0, "header": {}, "all_rows": []},
            )

        # 5. 获取 AI 配置（通过 context 注入，消除 packages→apps 反向依赖 T3-3）
        if context.ai_config_provider is None:
            raise AppError(
                code="ai_not_configured",
                message="AI 配置提供器未注入，无法获取大模型配置",
                retryable=False,
            )
        config: dict[str, Any] | None = await context.ai_config_provider()
        if config is None:
            raise AppError(
                code="ai_not_configured",
                message="AI 大模型未配置，请在平台治理 → AI 配置中开启",
                retryable=False,
            )

        # 6. 构建 LLM 请求（单条 user 消息，不分角色）
        if is_image_mode:
            # 多模态：图片 + 文本指令
            user_content: list[dict[str, Any]] = [
                {"type": "text", "text": f"{prompt}\n\n请根据以下图片内容提取数据。"},
            ]
            for img_data_url in content:
                user_content.append(
                    {
                        "type": "image_url",
                        "image_url": {"url": img_data_url},
                    }
                )
            messages = [
                {"role": "user", "content": user_content},
            ]
        else:
            # 纯文本模式
            user_message: str = f"{prompt}\n\n文件内容：\n{content}"
            messages = [
                {"role": "user", "content": user_message},
            ]

        request_body: dict[str, Any] = {
            "model": config["model_name"],
            "messages": messages,
            "chat_template_kwargs": {
                "enable_thinking": False,
            },
        }

        base_url: str = str(config["base_url"]).rstrip("/")
        url: str = f"{base_url}/chat/completions"
        headers: dict[str, str] = {
            "Authorization": f"Bearer {config['api_key']}",
            "Content-Type": "application/json",
        }

        try:
            async with httpx.AsyncClient(timeout=float(timeout), proxy=None) as client:
                resp = await client.post(url, headers=headers, json=request_body)
        except httpx.TimeoutException:
            raise AppError(
                code="ai_timeout",
                message=f"LLM 调用超时（{timeout} 秒）",
                retryable=True,
            ) from None
        except httpx.HTTPError as exc:
            # Server disconnected 等连接错误，自动重试一次
            err_msg = str(exc)[:200]
            if "disconnected" in err_msg.lower() or "remote" in err_msg.lower():
                import asyncio as _aio

                await _aio.sleep(2)
                try:
                    async with httpx.AsyncClient(
                        timeout=float(timeout + 120), proxy=None
                    ) as client2:
                        resp = await client2.post(url, headers=headers, json=request_body)
                except httpx.HTTPError as exc2:
                    raise AppError(
                        code="ai_request_failed",
                        message=f"LLM 请求失败（重试后）：{str(exc2)[:200]}",
                        retryable=True,
                    ) from None
            else:
                raise AppError(
                    code="ai_request_failed",
                    message=f"LLM 请求失败：{err_msg}",
                    retryable=True,
                ) from None

        if resp.status_code != 200:
            raise AppError(
                code="ai_request_failed",
                message=f"LLM API 返回 {resp.status_code}: {resp.text[:200]}",
                retryable=True,
            )

        # 6. 解析返回
        resp_data: dict[str, Any] = resp.json()
        choices: list[dict[str, Any]] = resp_data.get("choices", [])
        if not choices:
            raise AppError(
                code="ai_empty_response",
                message="LLM 返回空响应",
                retryable=True,
            )
        llm_content: str = choices[0]["message"]["content"]

        extracted_data: dict[str, Any] = _parse_llm_json(llm_content)
        header: dict[str, Any] = extracted_data.get("metadata", {})
        points: list[dict[str, Any]] = extracted_data.get("points", [])
        series: list[dict[str, Any]] = extracted_data.get("series", [])

        # 7. 从 points 推断列名（固定为 name, value, unit）
        if not points:
            columns: tuple[str, ...] = ()
            converted_rows: list[dict[str, Any]] = []
        else:
            columns = ("name", "value", "unit")
            converted_rows = points

        # 8. 构建 source_locations
        source_locs: list[dict[str, Any]] = [
            {"file": file_path.name, "row": idx} for idx in range(1, len(converted_rows) + 1)
        ]

        # 8. 构建 ObservationTable
        table = ObservationTable(
            columns=columns,
            rows=tuple(converted_rows),
            source_locations=tuple(source_locs),
        )

        return ComponentResult(
            outputs={"observations": table},
            summary=f"提取 {len(points)} 个指标，{len(series)} 组序列",
            metadata={
                "row_count": len(points),
                "header": header,
                "points": points,
                "series": series,
            },
        )

    @staticmethod
    async def _download_artifact_to_temp(
        context: ComponentContext,
        artifact_id_str: str,
    ) -> Path:
        """从 MinIO 下载工件到临时文件，返回临时文件路径。

        通过 context.artifact_service.get_bytes() 下载工件内容，
        写入临时文件，返回 Path 对象。临时文件由调用方负责清理。

        Args:
            context: 组件执行上下文（提供 artifact_service）。
            artifact_id_str: 工件 ID 字符串。

        Returns:
            Path: 临时文件路径。

        Raises:
            AppError: code="missing_dependency"，当 artifact_service 未注入。
            AppError: code="not_found"，当工件不存在。
        """
        artifact_service = context.artifact_service
        if artifact_service is None:
            raise AppError(
                code="missing_dependency",
                message="artifact_service 未注入，无法下载 artifact",
                retryable=False,
                fields={},
            )

        artifact_id = UUID(artifact_id_str)
        data: bytes = await artifact_service.get_bytes(artifact_id)

        # 写入临时文件
        fd, temp_path = tempfile.mkstemp(suffix=_guess_suffix(data))
        try:
            os.write(fd, data)
        finally:
            os.close(fd)

        return Path(temp_path)


def _extract_text(file_path: Path, engine: str = "auto", image_dpi: int = 200) -> str | list[str]:
    """从文件中提取文本内容。

    根据 engine 参数选择模式：
    - auto（默认）：根据文件后缀自动选择读取方式
      - .pdf: 先提取文字层，文字太少（平均每页 < 50 字符）自动切换到 image 模式
      - .txt/.md: 直接 UTF-8 读取
      - .jpg/.jpeg/.png: 返回 base64 图片列表（多模态）
      - .doc/.docx: 用 python-docx 读取
      - .xls/.xlsx: 用 openpyxl 读取
      - 其他: 尝试当文本读取
    - pymupdf: 使用 PyMuPDF 提取文字层（仅对 PDF 有效，向后兼容）
    - image: 使用 PyMuPDF 渲染为图片，返回 base64 列表（仅对 PDF 有效，向后兼容）
    - raw: 直接以 UTF-8 读取（向后兼容）

    Args:
        file_path: 文件路径。
        engine: 读取引擎（auto / pymupdf / image / raw），默认 auto。
        image_dpi: image 模式下的渲染分辨率。

    Returns:
        str | list[str]: 文本模式返回文本；image 模式返回 base64 图片列表。
    """
    suffix: str = file_path.suffix.lower()

    # 非 auto 模式：保留旧逻辑向后兼容
    if engine != "auto":
        if suffix != ".pdf":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if engine == "image":
            return _extract_pdf_as_images(file_path, image_dpi)
        if engine == "pymupdf":
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(str(file_path))
                text_parts: list[str] = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except ImportError:
                return file_path.read_text(encoding="utf-8", errors="ignore")
        return file_path.read_text(encoding="utf-8", errors="ignore")

    # auto 模式：按文件后缀自动分派
    if suffix == ".pdf":
        return _extract_pdf(file_path, image_dpi)
    elif suffix in (".txt", ".md", ""):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    elif suffix in (".jpg", ".jpeg", ".png"):
        return _extract_image_file(file_path)
    elif suffix in (".doc", ".docx"):
        return _extract_docx(file_path)
    elif suffix in (".xls", ".xlsx"):
        return _extract_xlsx(file_path)
    else:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(file_path: Path, image_dpi: int = 200) -> str | list[str]:
    """PDF 自动检测：先提取文字层，文字太少则切换到 image 模式。

    用 PyMuPDF 提取每页文字，计算平均每页字符数。如果平均每页 < 50 字符，
    判定为扫描版 PDF，自动切换到 image 模式（渲染图片发给多模态大模型识别）。

    Args:
        file_path: PDF 文件路径。
        image_dpi: image 模式下的渲染分辨率。

    Returns:
        str | list[str]: 文本模式返回文本；image 模式返回 base64 图片列表。
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(file_path))
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        full_text: str = "\n".join(text_parts)

        # 自动检测：如果文字层内容太少（平均每页 < 50 字符），判定为扫描版
        page_count: int = len(text_parts) if text_parts else 1
        avg_chars: float = len(full_text) / max(page_count, 1)
        if avg_chars < 50:
            # 扫描版 PDF，切换到 image 模式
            return _extract_pdf_as_images(file_path, image_dpi)
        return full_text
    except ImportError:
        # PyMuPDF 不可用，回退到直接读取
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf_as_images(file_path: Path, image_dpi: int = 200) -> list[str] | str:
    """将 PDF 渲染为 base64 图片列表（用于多模态大模型识别）。

    Args:
        file_path: PDF 文件路径。
        image_dpi: 渲染分辨率（DPI）。

    Returns:
        list[str] | str: base64 data URL 列表；PyMuPDF 不可用时回退到文本读取。
    """
    try:
        import base64

        import fitz  # PyMuPDF

        doc = fitz.open(str(file_path))
        images: list[str] = []
        for page in doc:
            pix = page.get_pixmap(dpi=image_dpi)
            img_bytes: bytes = pix.tobytes("png")
            b64: str = base64.b64encode(img_bytes).decode("utf-8")
            images.append(f"data:image/png;base64,{b64}")
        doc.close()
        return images
    except ImportError:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_image_file(file_path: Path) -> list[str]:
    """将图片文件转为 base64 data URL 列表（用于多模态大模型识别）。

    支持 .jpg/.jpeg/.png 格式。返回格式与 PDF image 模式一致，
    使下游 is_image_mode 判断无需修改。

    Args:
        file_path: 图片文件路径。

    Returns:
        list[str]: 包含单个 base64 data URL 的列表。
    """
    import base64

    suffix: str = file_path.suffix.lower()
    mime: str = "image/jpeg" if suffix in (".jpg", ".jpeg") else "image/png"
    with open(file_path, "rb") as f:
        b64: str = base64.b64encode(f.read()).decode("utf-8")
    return [f"data:{mime};base64,{b64}"]


def _extract_docx(file_path: Path) -> str:
    """提取 Word 文档文本。

    使用 python-docx 库读取 .doc/.docx 文件的段落文本。

    Args:
        file_path: Word 文件路径。

    Returns:
        str: 文档中所有段落的文本，以换行连接。

    Raises:
        AppError: code="validation_failed"，当 python-docx 依赖不可用时。
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []
        # 提取段落文本
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # 提取表格文本（按行按单元格）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Word 文件需要 python-docx 依赖，请执行 pip install python-docx",
            retryable=False,
        ) from None


def _extract_xlsx(file_path: Path) -> str:
    """提取 Excel 文件文本。

    使用 openpyxl 库读取 .xls/.xlsx 文件，将每个工作表的每行以 tab 分隔，
    行间以换行连接。

    Args:
        file_path: Excel 文件路径。

    Returns:
        str: 所有工作表的文本内容。

    Raises:
        AppError: code="validation_failed"，当 openpyxl 依赖不可用时。
    """
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells: list[str] = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
        wb.close()
        return "\n".join(lines)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Excel 文件需要 openpyxl 依赖，请执行 pip install openpyxl",
            retryable=False,
        ) from None


def _parse_llm_json(content: str) -> dict[str, Any]:
    """从 LLM 返回内容中提取 JSON 对象（3 级 fallback）。"""
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass

    pattern: str = r"```(?:json)?\s*\n?(.*?)\n?```"
    match = re.search(pattern, content, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass

    start: int = content.find("{")
    end: int = content.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(content[start : end + 1])
        except json.JSONDecodeError:
            pass

    raise AppError(
        code="ai_parse_failed",
        message=f"无法从 LLM 响应中解析 JSON：{content[:200]}",
        retryable=True,
    )


def _guess_suffix(data: bytes) -> str:
    """根据文件内容魔数推断文件后缀。

    通过检查文件头部的魔数字节判断文件类型，返回合适的后缀名。
    支持识别 PDF、PNG、JPEG、Office Open XML（xlsx/docx）格式。
    无法识别时返回空字符串（由系统分配默认后缀）。

    Args:
        data: 文件内容字节。

    Returns:
        str: 文件后缀（如 ``".pdf"``），无法识别返回 ``""``。
    """
    if data[:4] == b"%PDF":
        return ".pdf"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return ".png"
    if data[:2] == b"\xff\xd8":
        return ".jpg"
    # Office Open XML (xlsx/docx) — ZIP 格式，魔数 PK\x03\x04
    # 通过检查 ZIP 内部目录区分 docx（含 word/）和 xlsx（含 xl/）
    if data[:4] == b"PK\x03\x04":
        try:
            import zipfile

            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = zf.namelist()
                if any(n.startswith("word/") for n in names):
                    return ".docx"
                if any(n.startswith("xl/") for n in names):
                    return ".xlsx"
        except Exception:
            pass
        return ".xlsx"  # 默认当 xlsx
    return ""
