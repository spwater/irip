"""文本提取公共模块。

从 ``llm_converter/converter.py`` 迁移而来，提供统一的文件文本提取能力，
供 ``llm_converter`` 与 ``component_preview.py`` 共享。

主入口 ``extract_text`` 按文件后缀与引擎策略分发到各子函数：
- PDF: 优先提取文字层，文字过少则用 PaddleOCR 识别（返回纯文本）
- 图片 (.jpg/.jpeg/.png): 用 PaddleOCR 提取文字（返回纯文本）
- Word (.doc/.docx): 提取段落 + 表格文本
- Excel (.xls/.xlsx): 提取全部单元格文本
- 其他/文本: 直接读取 UTF-8 文本

注意：所有提取均返回 str（纯文本），不再返回 list[str]（base64 图片）。
图片文件通过 PaddleOCR 转为文本，避免依赖多模态 LLM。
"""

import logging
from pathlib import Path
from typing import Any

from packages.common.errors import AppError

logger = logging.getLogger(__name__)


def extract_text(file_path: Path, engine: str = "auto", image_dpi: int = 200) -> str:
    """从文件中提取文本内容。

    所有文件类型均返回 str（纯文本）。图片和扫描件 PDF 通过 PaddleOCR
    提取文字，不依赖多模态 LLM。

    Args:
        file_path: 文件路径。
        engine: 提取引擎（auto / pymupdf / raw），默认 auto。
            - auto: 按后缀自动选择最佳策略。
            - pymupdf: 仅对 PDF 生效，提取文字层。
            - raw: 直接读取文本内容。
        image_dpi: PDF 转图片的 DPI（OCR 模式时使用），默认 200。

    Returns:
        str: 文本内容。
    """
    suffix: str = file_path.suffix.lower()

    if engine != "auto":
        if suffix != ".pdf":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if engine == "pymupdf":
            try:
                import fitz

                doc = fitz.open(str(file_path))
                text_parts: list[str] = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except ImportError:
                return file_path.read_text(encoding="utf-8", errors="ignore")
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        return _extract_pdf(file_path, image_dpi)
    elif suffix in (".txt", ".md", ""):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    elif suffix in (".jpg", ".jpeg", ".png"):
        return _extract_image_with_ocr(file_path)
    elif suffix in (".doc", ".docx"):
        return _extract_docx(file_path)
    elif suffix in (".xls", ".xlsx"):
        return _extract_xlsx(file_path)
    else:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(file_path: Path, image_dpi: int = 200) -> str:
    """PDF 自动检测：先提取文字层，文字太少则用 PaddleOCR 识别。"""
    try:
        import fitz

        doc = fitz.open(str(file_path))
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        full_text: str = "\n".join(text_parts)

        page_count: int = len(text_parts) if text_parts else 1
        avg_chars: float = len(full_text) / max(page_count, 1)
        if avg_chars < 50:
            # 文字层不足（扫描件）→ 用 PaddleOCR 提取文字
            logger.info("PDF 文字层不足 (avg_chars=%.1f)，使用 PaddleOCR", avg_chars)
            return _extract_pdf_with_ocr(file_path, image_dpi)
        return full_text
    except ImportError:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf_with_ocr(file_path: Path, image_dpi: int = 200) -> str:
    """用 PaddleOCR 对 PDF 每页做 OCR，返回纯文本。"""
    try:
        import os
        import tempfile

        import fitz
    except ImportError:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    # 渲染每页为临时图片
    doc = fitz.open(str(file_path))
    tmp_paths: list[str] = []
    try:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=image_dpi)
            fd, tmp = tempfile.mkstemp(suffix=".png", prefix=f"pdf_ocr_{i}_")
            try:
                with open(tmp, "wb") as f:
                    f.write(pix.tobytes("png"))
            except Exception:
                os.unlink(tmp)
                raise
            tmp_paths.append(tmp)
    finally:
        doc.close()

    # 对每页做 OCR
    all_text: list[str] = []
    try:
        for tmp in tmp_paths:
            ocr_text = _extract_image_with_ocr(Path(tmp))
            if ocr_text.strip():
                all_text.append(ocr_text)
    finally:
        for tmp in tmp_paths:
            try:
                os.unlink(tmp)
            except Exception:
                logger.debug("cleanup failed", exc_info=True)

    return "\n\n".join(all_text)


#: PaddleOCR 全局单例（首次调用时初始化，后续复用，避免重复加载模型）。
_ocr_engine = None


def _get_ocr_engine() -> Any:
    """获取或初始化 PaddleOCR 单例。"""
    global _ocr_engine
    if _ocr_engine is not None:
        return _ocr_engine
    try:
        from paddleocr import PaddleOCR
    except ImportError:
        logger.warning("PaddleOCR 未安装，图片文字提取返回空")
        return None
    _ocr_engine = PaddleOCR(
        lang="ch",
        text_detection_model_name="PP-OCRv6_small_det",
        text_recognition_model_name="PP-OCRv6_small_rec",
    )
    return _ocr_engine


def _extract_image_with_ocr(file_path: Path) -> str:
    """用 PaddleOCR 从图片提取文字。

    如果 PaddleOCR 未安装，返回空字符串（LLM 会收到空内容，返回空结果）。
    """
    ocr = _get_ocr_engine()
    if ocr is None:
        return ""

    try:
        result = ocr.predict(str(file_path))

        # result 是 list[dict]，每个 dict 含 rec_texts (list[str])
        lines: list[str] = []
        if result:
            for page in result:
                if isinstance(page, dict):
                    rec_texts = page.get("rec_texts", [])
                    lines.extend(rec_texts)
        return "\n".join(lines)
    except Exception as exc:
        logger.warning("PaddleOCR 识别失败: %s", exc)
        return ""


def _extract_docx(file_path: Path) -> str:
    """提取 Word 文档文本。"""
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Word 文件需要 python-docx 依赖",
            retryable=False,
        ) from None


def _extract_xlsx(file_path: Path) -> str:
    """提取 Excel 文件文本。"""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells: list[str] = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
        wb.close()
        return "\n".join(lines)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Excel 文件需要 openpyxl 依赖",
            retryable=False,
        ) from None
