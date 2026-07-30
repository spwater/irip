"""大模型解析器插件。

从文件中提取文本，调用大模型转为结构化数据 {metadata, points, series}。
支持 PDF、图片、Word、Excel、纯文本等多种文件格式。

从 ez_scan_extractor.py 迁移 LLM 解析逻辑，使其与 XRD 解析器
走同一插件调用路径。
"""

import asyncio
import json
import re
from pathlib import Path
from typing import Any

import httpx

from packages.common.errors import AppError
from packages.common.safe_http import SafeHTTPClient
from packages.plugins.protocol import ConverterResult


class LlmConverter:
    """大模型解析器。

    从文件提取文本 → 调用 LLM → 返回结构化数据。
    """

    async def execute(self, params: dict[str, Any]) -> dict[str, Any]:
        """提取文件文本 → 调用大模型 → 返回结构化数据。

        Args:
            params: 参数字典，包含:
                - file_path: 文件路径（必填）
                - prompt: LLM 提示词（必填）
                - file_engine: 文件读取方式（auto/pymupdf/image/raw），默认 auto
                - ai_config: AI 配置字典（base_url/api_key/model_name）
                - timeout: 超时秒数，默认 300

        Returns:
            dict: ``{"metadata": {...}, "points": [...], "series": [...]}``
        """
        file_path = Path(params["file_path"])
        prompt: str = params.get("prompt", "")
        engine: str = params.get("file_engine", "auto")
        image_dpi: int = params.get("image_dpi", 200)
        max_chars: int = params.get("max_content_chars", 999999999)
        timeout: int = params.get("timeout", 300)
        config: dict[str, Any] | None = params.get("ai_config")

        if not prompt:
            raise AppError(
                code="validation_failed",
                message="缺少 prompt 参数",
                retryable=False,
            )

        if config is None:
            raise AppError(
                code="ai_not_configured",
                message="AI 大模型未配置，请在平台治理 → AI 配置中开启",
                retryable=False,
            )

        # 1. 提取文件文本
        content = await asyncio.to_thread(_extract_text, file_path, engine, image_dpi)
        if isinstance(content, str) and len(content) > max_chars:
            content = content[:max_chars]

        # 2. 空内容直接返回空结果
        is_image_mode = isinstance(content, list)
        if (is_image_mode and len(content) == 0) or (not is_image_mode and not content.strip()):
            return ConverterResult().to_dict()

        # 3. 构建 LLM 请求
        if is_image_mode:
            user_content: list[dict[str, Any]] = [
                {"type": "text", "text": f"{prompt}\n\n请根据以下图片内容提取数据。"},
            ]
            for img_data_url in content:
                user_content.append(
                    {
                        "type": "image_url",
                        "image_url": {"url": img_data_url},
                    }
                )
            messages = [{"role": "user", "content": user_content}]
        else:
            user_message = f"{prompt}\n\n文件内容：\n{content}"
            messages = [{"role": "user", "content": user_message}]

        request_body: dict[str, Any] = {
            "model": config["model_name"],
            "messages": messages,
            "chat_template_kwargs": {"enable_thinking": False},
            "temperature": 0.0,
            "seed": 42,
        }

        base_url: str = str(config["base_url"]).rstrip("/")
        url: str = f"{base_url}/chat/completions"
        headers: dict[str, str] = {
            "Authorization": f"Bearer {config['api_key']}",
            "Content-Type": "application/json",
        }

        # 4. 调用 LLM
        resp = await _call_llm(url, headers, request_body, timeout)

        # 5. 解析返回
        resp_data: dict[str, Any] = resp.json()
        choices: list[dict[str, Any]] = resp_data.get("choices", [])
        if not choices:
            raise AppError(
                code="ai_empty_response",
                message="LLM 返回空响应",
                retryable=True,
            )
        llm_content: str = choices[0]["message"]["content"]

        extracted_data: dict[str, Any] = _parse_llm_json(llm_content)
        return ConverterResult(
            metadata=extracted_data.get("metadata", {}),
            points=extracted_data.get("points", []),
            series=extracted_data.get("series", []),
        ).to_dict()


async def _call_llm(
    url: str,
    headers: dict[str, str],
    body: dict[str, Any],
    timeout: int,  # noqa: ASYNC109
) -> httpx.Response:
    """调用 LLM API，含断线重试（H-05: 使用 SafeHTTPClient）。"""
    try:
        # H-05: 使用 SafeHTTPClient（SSRF 防护 + 流式大小限制）
        async with SafeHTTPClient(timeout=float(timeout), max_size=10 * 1024 * 1024) as client:
            resp = await client.post(url, headers=headers, json=body)
    except httpx.TimeoutException:
        raise AppError(
            code="ai_timeout",
            message=f"LLM 调用超时（{timeout} 秒）",
            retryable=True,
        ) from None
    except httpx.HTTPError as exc:
        err_msg = str(exc)[:200]
        if "disconnected" in err_msg.lower() or "remote" in err_msg.lower():
            await asyncio.sleep(2)
            try:
                async with SafeHTTPClient(
                    timeout=float(timeout + 120), max_size=10 * 1024 * 1024
                ) as client2:
                    resp = await client2.post(url, headers=headers, json=body)
            except httpx.HTTPError as exc2:
                raise AppError(
                    code="ai_request_failed",
                    message=f"LLM 请求失败（重试后）：{str(exc2)[:200]}",
                    retryable=True,
                ) from None
        else:
            raise AppError(
                code="ai_request_failed",
                message=f"LLM 请求失败：{err_msg}",
                retryable=True,
            ) from None

    if resp.status_code != 200:
        raise AppError(
            code="ai_request_failed",
            message=f"LLM API 返回 {resp.status_code}: {resp.text[:200]}",
            retryable=True,
        )
    return resp


def _extract_text(file_path: Path, engine: str = "auto", image_dpi: int = 200) -> str | list[str]:
    """从文件中提取文本内容。"""
    suffix: str = file_path.suffix.lower()

    if engine != "auto":
        if suffix != ".pdf":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if engine == "image":
            return _extract_pdf_as_images(file_path, image_dpi)
        if engine == "pymupdf":
            try:
                import fitz

                doc = fitz.open(str(file_path))
                text_parts: list[str] = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except ImportError:
                return file_path.read_text(encoding="utf-8", errors="ignore")
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        return _extract_pdf(file_path, image_dpi)
    elif suffix in (".txt", ".md", ""):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    elif suffix in (".jpg", ".jpeg", ".png"):
        return _extract_image_file(file_path)
    elif suffix in (".doc", ".docx"):
        return _extract_docx(file_path)
    elif suffix in (".xls", ".xlsx"):
        return _extract_xlsx(file_path)
    else:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(file_path: Path, image_dpi: int = 200) -> str | list[str]:
    """PDF 自动检测：先提取文字层，文字太少则切换到 image 模式。"""
    try:
        import fitz

        doc = fitz.open(str(file_path))
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        full_text: str = "\n".join(text_parts)

        page_count: int = len(text_parts) if text_parts else 1
        avg_chars: float = len(full_text) / max(page_count, 1)
        if avg_chars < 50:
            return _extract_pdf_as_images(file_path, image_dpi)
        return full_text
    except ImportError:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf_as_images(file_path: Path, image_dpi: int = 200) -> list[str] | str:
    """将 PDF 渲染为 base64 图片列表。"""
    try:
        import base64

        import fitz

        doc = fitz.open(str(file_path))
        images: list[str] = []
        for page in doc:
            pix = page.get_pixmap(dpi=image_dpi)
            img_bytes: bytes = pix.tobytes("png")
            b64: str = base64.b64encode(img_bytes).decode("utf-8")
            images.append(f"data:image/png;base64,{b64}")
        doc.close()
        return images
    except ImportError:
        return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_image_file(file_path: Path) -> list[str]:
    """将图片文件转为 base64 data URL 列表。"""
    import base64

    suffix: str = file_path.suffix.lower()
    mime: str = "image/jpeg" if suffix in (".jpg", ".jpeg") else "image/png"
    with open(file_path, "rb") as f:
        b64: str = base64.b64encode(f.read()).decode("utf-8")
    return [f"data:{mime};base64,{b64}"]


def _extract_docx(file_path: Path) -> str:
    """提取 Word 文档文本。"""
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Word 文件需要 python-docx 依赖",
            retryable=False,
        ) from None


def _extract_xlsx(file_path: Path) -> str:
    """提取 Excel 文件文本。"""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells: list[str] = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
        wb.close()
        return "\n".join(lines)
    except ImportError:
        raise AppError(
            code="validation_failed",
            message="读取 Excel 文件需要 openpyxl 依赖",
            retryable=False,
        ) from None


def _parse_llm_json(content: str) -> dict[str, Any]:
    """从 LLM 返回内容中提取 JSON 对象（3 级 fallback）。"""
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass

    pattern: str = r"```(?:json)?\s*\n?(.*?)\n?```"
    match = re.search(pattern, content, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass

    start: int = content.find("{")
    end: int = content.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(content[start : end + 1])
        except json.JSONDecodeError:
            pass

    raise AppError(
        code="ai_parse_failed",
        message=f"无法从 LLM 响应中解析 JSON：{content[:200]}",
        retryable=True,
    )
